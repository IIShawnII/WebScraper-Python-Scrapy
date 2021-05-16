from docx import Document


def saveToDocx(data, fileName):

    document = Document()

    document.add_heading("Email: " + data["email"])
    document.add_paragraph("UUID: " + data["uuid"])
    document.add_heading("Page Url:" + data["url"], level=2)

    document.add_heading("Title", level=3)
    document.add_paragraph(data["title"][0])

    document.add_heading("HyperLink's", level=3)
    for item in data["link"]:
        document.add_paragraph(item)

    document.add_heading("Span's", level=3)
    for item in data["span"]:
        document.add_paragraph(item)

    document.add_heading("Paragraph's", level=3)
    for item in data["paragraph"]:
        document.add_paragraph(item)

    document.add_heading("Heading's 1", level=3)
    for item in data["heading 1"]:
        document.add_paragraph(item)

    document.add_heading("Heading's 2", level=3)
    for item in data["heading 2"]:
        document.add_paragraph(item)

    document.add_heading("Heading's 3", level=3)
    for item in data["heading 3"]:
        document.add_paragraph(item)

    document.add_heading("Heading's 4", level=3)
    for item in data["heading 4"]:
        document.add_paragraph(item)

    document.add_heading("Other's", level=3)
    for item in data["other"]:
        document.add_paragraph(item)

    document.save(fileName+".docx")
